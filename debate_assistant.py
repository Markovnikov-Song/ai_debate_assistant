"""
多智能体辩论决策助手
- 多个 AI 角色从不同立场辩论，帮助用户做决策
- 支持用户插嘴、自由辩、学术资料搜索、历史记录
- 数据持久化到 GitHub 私有仓库
"""

import io
import re
import json
import os
import random
import requests
from datetime import datetime

import streamlit as st
from openai import OpenAI
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv

from auth import render_auth_page, render_admin_panel, render_admin_page, logout
from github_storage import (
    load_agent_config as gh_load_agent_config,
    save_agent_config as gh_save_agent_config,
    load_debate, save_debate, delete_debate, list_debates,
    load_feedback, save_feedback as gh_save_feedback,
)

load_dotenv()

# ── 登录拦截（未登录直接停止渲染） ────────────────────────────────────
if not render_auth_page():
    st.stop()

# ── 基础配置 ──────────────────────────────────────────────────────────
def _cfg(key, default=""):
    try:
        return st.secrets[key]
    except Exception:
        return os.getenv(key, default)

API_KEY      = _cfg("API_KEY",  "your-api-key-here")
BASE_URL     = _cfg("BASE_URL", "https://api.siliconflow.cn/v1")
MODEL        = _cfg("MODEL",    "deepseek-ai/DeepSeek-V3")
CURRENT_USER = st.session_state.get("current_user", "guest")

client = OpenAI(api_key=API_KEY, base_url=BASE_URL)

# ── 页面配置 ──────────────────────────────────────────────────────────
st.set_page_config(page_title="多智能体辩论助手", page_icon="🎯", layout="centered")
st.markdown("""
<style>
@media (max-width: 768px) {
    .block-container { padding: 1rem 0.75rem; padding-top: 2.5rem !important; }
    h1 { font-size: 1.4rem !important; word-break: break-all; }
    h2, h3 { font-size: 1rem !important; }
    p, li { font-size: 0.95rem !important; line-height: 1.6; }
    .stButton > button { width: 100%; }
}
/* 修复新版 Streamlit download_button 大图标卡片样式 */
[data-testid="stDownloadButton"] > button {
    display: flex !important;
    flex-direction: row !important;
    align-items: center !important;
    justify-content: center !important;
    padding: 0.4rem 0.75rem !important;
    font-size: 0.9rem !important;
    height: auto !important;
    min-height: unset !important;
}
[data-testid="stDownloadButton"] > button > span {
    display: flex !important;
    flex-direction: row !important;
    align-items: center !important;
    gap: 0.3rem !important;
}
[data-testid="stDownloadButton"] svg {
    display: none !important;
}
</style>
""", unsafe_allow_html=True)

# ── Session State 初始化 ──────────────────────────────────────────────
_DEFAULTS = {
    "topic": "",
    "user_context": "",
    "debate_history": [],
    "debate_round": 0,
    "current_history_id": None,
    "custom_agents": [],
    "interrupt_rounds": 3,
    "last_summary": "",
    "pending_user_speech": None,   # 待处理的用户插嘴 {target, content, round}
    "search_results": {},          # 学术搜索缓存 {"round_agentName": "..."}
    "enable_search": True,
    "is_debating": False,          # 防止生成中重复点击
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

# ════════════════════════════════════════════════════════════════════
# 角色配置
# ════════════════════════════════════════════════════════════════════

_DEFAULT_AGENTS = [
    {"name": "支持派",     "prompt": "你是立场鲜明的支持派，说话自然、有逻辑、有情绪，但不偏激。先承接上一轮内容，针对性反驳对方，再讲自己的观点。核心理由、关键结论必须加粗。"},
    {"name": "反对派",     "prompt": "你是清醒理性的反对派，指出问题一针见血。先针对性反驳对方上一轮观点，再提出自己的看法。核心理由、风险、弊端必须加粗。"},
    {"name": "中立理性派", "prompt": "你是中立理性派，不站队、只讲事实。指出双方合理与不合理的地方，语气温和客观。关键判断、核心差异必须加粗。"},
    {"name": "长期视角派", "prompt": "你是长期视角观察者，站在未来发展和长远影响的角度分析，像过来人给建议。结合长远收益和潜在风险点评当前辩论。长期影响、未来收益必须加粗。"},
]

def load_agent_config() -> list:
    data = gh_load_agent_config(CURRENT_USER)
    return data if data else _DEFAULT_AGENTS

def save_agent_config(agents: list):
    gh_save_agent_config(CURRENT_USER, agents)

def auto_generate_agents(topic: str) -> list:
    """让 AI 根据议题自动生成立场角色"""
    prompt = f"""辩论议题：{topic}

请为这个议题生成4-5个立场鲜明的辩论角色，每个角色代表一种具体观点或路线。

输出格式（严格按此格式，每个角色一块）：
角色名: xxx
人设: xxx（50字以内，说明立场、说话风格、核心论点方向，结尾加"核心论点必须加粗"）

---"""
    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7, max_tokens=600,
        ).choices[0].message.content.strip()
        agents = []
        for block in resp.split("---"):
            name, persona = "", ""
            for line in block.strip().splitlines():
                if line.startswith("角色名:"):
                    name = line.split(":", 1)[-1].strip()
                elif line.startswith("人设:"):
                    persona = line.split(":", 1)[-1].strip()
            if name and persona:
                agents.append({"name": name, "prompt": persona})
        return agents
    except Exception:
        return []

if not st.session_state.custom_agents:
    st.session_state.custom_agents = load_agent_config()

# ════════════════════════════════════════════════════════════════════
# Prompt 模板
# ════════════════════════════════════════════════════════════════════

_RULES_FIRST_ROUND = """
【辩论规则】
这是第一轮，还没有其他人发言。请直接从你的立场出发，陈述你对这个议题最核心的观点。
1. 不要假设或捏造其他人的观点。
2. 核心理由、关键结论必须用Markdown加粗。
3. 语言自然、像真人，不要机器腔。
4. 控制在150字以内，简明扼要。
"""

_RULES_NORMAL = """
【辩论规则】
1. 先直接回应上一轮的内容，必须反驳或承接，不能自说自话。
2. 可以质疑对方的逻辑、数据、前提，像真实辩论一样。
3. 核心理由、关键结论、反驳点必须用Markdown加粗。
4. 语言自然、像真人，不要机器腔。
5. 控制在150字以内，简明扼要，不要废话。
"""

_SUMMARIZER_PROMPT = """
你是专业的最终决策总结者，基于完整辩论历史输出清晰可落地的总结。
要求：
1. **核心争议点**：提炼2-3个最核心的分歧。
2. **各方核心观点**：简要总结每个角色的核心立场。
3. **最终综合建议**：给出明确的决策方向（支持/反对/折中），贴合用户的具体场景。
4. **决策依据**：列出3-5条支持该建议的核心论据。
格式：用中文数字分点，最重要的结论必须加粗，语言简洁。
"""

def get_debate_prompt(agent_prompt: str, is_first_round: bool = False) -> str:
    return agent_prompt + (_RULES_FIRST_ROUND if is_first_round else _RULES_NORMAL)

def get_interrupt_prompt(agent_prompt: str, target_name: str, target_content: str) -> str:
    return agent_prompt + f"""
【自由辩规则】
现在是自由辩环节，你要打断并反驳"{target_name}"刚才说的这段话：
"{target_content}"
要求：
1. 直接针对上面这段话的漏洞、矛盾或不合理之处发起反驳或追问
2. 不需要完整阐述自己的立场，只需要一针见血地指出问题
3. 语气可以尖锐，像真实辩论赛的自由辩环节
4. 控制在100字以内，简短有力
"""

def get_user_speech_prompt(agent_prompt: str, agent_name: str, target_name: str,
                            user_content: str, is_targeted: bool) -> str:
    """用户插嘴后各角色的回应 prompt"""
    if is_targeted:
        return agent_prompt + f"""
【本轮特殊规则】
用户直接向你（{agent_name}）提出了质疑：
"{user_content}"
要求：
1. 必须先正面回应用户的质疑，不能回避
2. 回应完后可以继续阐述自己的立场
3. 核心反驳点必须加粗
4. 控制在150字以内
"""
    return agent_prompt + f"""
【本轮背景】
用户刚才向{target_name}提出了质疑："{user_content}"
你可以借用户的质疑支持或补充自己的立场，也可以反驳用户的质疑。
核心观点必须加粗，控制在150字以内。
"""

# ════════════════════════════════════════════════════════════════════
# 学术搜索
# ════════════════════════════════════════════════════════════════════

def _search_semantic_scholar(query: str, max_results: int = 3) -> list[dict]:
    """Semantic Scholar 学术论文搜索（免费，无需 key）"""
    try:
        r = requests.get(
            "https://api.semanticscholar.org/graph/v1/paper/search",
            params={"query": query, "limit": max_results,
                    "fields": "title,abstract,year,externalIds"},
            timeout=8,
        )
        if not r.ok:
            return []
        results = []
        for item in r.json().get("data", []):
            abstract = item.get("abstract") or ""
            if not abstract:
                continue
            doi = (item.get("externalIds") or {}).get("DOI", "")
            url = f"https://doi.org/{doi}" if doi else \
                  f"https://www.semanticscholar.org/paper/{item.get('paperId','')}"
            results.append({
                "source": "📄 学术论文",
                "title": item.get("title", ""),
                "year":  item.get("year", ""),
                "text":  abstract[:300],
                "url":   url,
            })
        return results
    except Exception:
        return []

def _search_wikipedia(query: str, max_results: int = 2) -> list[dict]:
    """中文维基百科搜索"""
    try:
        r = requests.get(
            "https://zh.wikipedia.org/w/api.php",
            params={"action": "query", "list": "search", "srsearch": query,
                    "srlimit": max_results, "format": "json", "uselang": "zh"},
            timeout=8,
        )
        if not r.ok:
            return []
        results = []
        for item in r.json().get("query", {}).get("search", []):
            snippet = re.sub(r'<[^>]+>', '', item.get("snippet", ""))
            if not snippet:
                continue
            title = item.get("title", "")
            results.append({
                "source": "📖 维基百科",
                "title": title,
                "year":  "",
                "text":  snippet[:300],
                "url":   f"https://zh.wikipedia.org/wiki/{requests.utils.quote(title)}",
            })
        return results
    except Exception:
        return []

def search_for_agent(topic: str, agent_name: str, agent_prompt: str) -> str:
    """为角色生成搜索关键词并检索学术资料，返回带来源标注的 Markdown 文本"""
    try:
        # 让 AI 生成针对该角色立场的搜索关键词
        kw_resp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": f"""议题：{topic}
你的角色立场：{agent_prompt[:100]}
请生成2个英文搜索关键词（用于学术搜索）和1个中文关键词（用于百科搜索），支持你的立场。
格式：
英文1: xxx
英文2: xxx
中文: xxx"""}],
            temperature=0.3, max_tokens=80,
        ).choices[0].message.content.strip()

        en_kws, zh_kw = [], topic
        for line in kw_resp.splitlines():
            if line.startswith("英文"):
                kw = line.split(":", 1)[-1].strip()
                if kw:
                    en_kws.append(kw)
            elif line.startswith("中文"):
                zh_kw = line.split(":", 1)[-1].strip() or topic

        all_results = []
        for kw in en_kws[:2]:
            all_results.extend(_search_semantic_scholar(kw, max_results=2))
        all_results.extend(_search_wikipedia(zh_kw, max_results=2))

        if not all_results:
            return ""

        lines = []
        for item in all_results[:4]:
            year = f"（{item['year']}）" if item["year"] else ""
            lines.append(f"{item['source']} [{item['title']}{year}]({item['url']})\n  {item['text']}")
        return "\n\n".join(lines)
    except Exception:
        return ""

# ════════════════════════════════════════════════════════════════════
# AI 调用
# ════════════════════════════════════════════════════════════════════

def format_history_recent(hist: list, last_n_rounds: int = 2) -> str:
    """取最近 N 轮辩论内容作为上下文"""
    if not hist:
        return ""
    max_r  = max(x["round"] for x in hist)
    cutoff = max(1, max_r - last_n_rounds + 1)
    lines  = []
    for x in hist:
        if x["round"] < cutoff:
            continue
        if x["type"] == "agent_speech":
            lines.append(f"第{x['round']}轮 - {x['name']}：{x['content']}\n\n")
        elif x["type"] == "user_speech":
            lines.append(f"第{x['round']}轮 - 【用户插嘴{x.get('target','')}】：{x['content']}\n\n")
    return "".join(lines)

def format_history_full(hist: list) -> str:
    """完整辩论历史，用于生成总结"""
    lines = []
    for x in hist:
        if x["type"] == "agent_speech":
            lines.append(f"第{x['round']}轮 - {x['name']}：{x['content']}\n\n")
        elif x["type"] == "user_speech":
            lines.append(f"第{x['round']}轮 - 【用户插嘴{x.get('target','')}】：{x['content']}\n\n")
    return "".join(lines)

def call_llm(prompt: str, history_context: str = "", max_tokens: int = 400,
             round_num=None, agent_name: str = None) -> str:
    """调用 LLM，自动注入议题、搜索资料和历史上下文"""
    context = f"核心辩论议题：{st.session_state.topic}\n"
    if st.session_state.user_context:
        context += f"用户补充条件：{st.session_state.user_context}\n"

    # 注入该角色的学术搜索结果
    search_data = st.session_state.get("search_results", {})
    search_key  = f"{round_num}_{agent_name}" if agent_name else round_num
    if search_key and search_data.get(search_key):
        context += f"【本轮学术参考资料（请引用来支撑你的论点）】\n{search_data[search_key]}\n\n"

    if history_context:
        context += f"近期辩论内容：\n{history_context}\n"
    context += f"你的任务：{prompt}"

    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": context}],
            temperature=0.8,
            max_tokens=max_tokens,
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"AI调用失败：{str(e)}")
        return "调用失败"

# ════════════════════════════════════════════════════════════════════
# 历史记录存储
# ════════════════════════════════════════════════════════════════════

def _build_debate_data(filename: str) -> dict:
    """构建辩论记录数据结构"""
    return {
        "id":             filename,
        "topic":          st.session_state.topic,
        "user_context":   st.session_state.user_context,
        "debate_round":   st.session_state.debate_round,
        "create_time":    datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "debate_content": st.session_state.debate_history,
        "agents_used":    st.session_state.custom_agents,
    }

def auto_save():
    """每轮结束后静默保存，同一议题覆盖上次自动存档"""
    if not st.session_state.topic or not st.session_state.debate_history:
        return
    topic_key = st.session_state.topic[:20].replace(" ", "_").replace("？", "_").replace("?", "_")
    filename  = f"auto_{topic_key}.json"
    save_debate(CURRENT_USER, filename, _build_debate_data(filename))

def save_debate_history() -> str | None:
    """手动保存，生成带时间戳的唯一文件名"""
    if not st.session_state.topic or not st.session_state.debate_history:
        return None
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    topic_key = st.session_state.topic[:12].replace(" ", "_").replace("？", "_").replace("?", "_")
    filename  = f"{timestamp}_{topic_key}.json"
    save_debate(CURRENT_USER, filename, _build_debate_data(filename))
    return filename

def load_debate_history(filename: str):
    """从 GitHub 加载辩论记录到 session state"""
    data = load_debate(CURRENT_USER, filename)
    if not data:
        st.error("记录不存在")
        return
    st.session_state.topic              = data["topic"]
    st.session_state.user_context       = data["user_context"]
    st.session_state.debate_round       = data["debate_round"]
    st.session_state.debate_history     = data["debate_content"]
    st.session_state.current_history_id = filename
    st.session_state.last_summary       = ""
    if "agents_used" in data:
        st.session_state.custom_agents = data["agents_used"]
    st.rerun()

def get_all_history() -> list[str]:
    files = list_debates(CURRENT_USER)
    files.sort(reverse=True)
    return files

def get_history_label(filename: str) -> str:
    try:
        data   = load_debate(CURRENT_USER, filename)
        prefix = "🔄 " if filename.startswith("auto_") else ""
        return f"{prefix}{data.get('topic','未知议题')[:16]}  [{data.get('create_time','')[:16]}]"
    except Exception:
        return filename

def delete_history(filename: str):
    delete_debate(CURRENT_USER, filename)
    if st.session_state.current_history_id == filename:
        for k, v in {"topic": "", "user_context": "", "debate_history": [],
                     "debate_round": 0, "current_history_id": None}.items():
            st.session_state[k] = v
    st.rerun()

# ════════════════════════════════════════════════════════════════════
# 导出功能
# ════════════════════════════════════════════════════════════════════

def _md_to_html(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
    return text.replace('\n', '<br>')

def _add_markdown_paragraph(doc, prefix_bold: str, text: str):
    """将含 **粗体** 的 Markdown 文本写入 Word 段落，正确渲染粗体"""
    p = doc.add_paragraph()
    if prefix_bold:
        p.add_run(prefix_bold).bold = True
    for i, part in enumerate(re.split(r'\*\*(.+?)\*\*', text)):
        if part:
            p.add_run(part).bold = (i % 2 == 1)
    return p

def export_to_html() -> tuple:
    if not st.session_state.topic or not st.session_state.debate_history:
        return None, None
    rows = []
    for r in range(1, st.session_state.debate_round + 1):
        rows.append(f'<h2>第{r}轮</h2>')
        for item in [x for x in st.session_state.debate_history if x["round"] == r]:
            if item["type"] == "user_speech":
                rows.append(f'<div class="user-speech"><strong>【用户插嘴 → {item.get("target","")}】</strong><br>{_md_to_html(item["content"])}</div>')
            else:
                rows.append(f'<div class="speech"><div class="agent-name">{item["name"]}</div><div class="agent-content">{_md_to_html(item["content"])}</div></div>')
        rows.append('<hr>')
    summary_html = f'<h2>决策总结</h2><div class="summary">{_md_to_html(st.session_state.last_summary)}</div>' \
                   if st.session_state.last_summary else ""
    html = f"""<!DOCTYPE html>
<html lang="zh">
<head>
<meta charset="UTF-8">
<title>辩论记录：{st.session_state.topic}</title>
<style>
  body {{ font-family: "PingFang SC", "Microsoft YaHei", sans-serif; max-width: 800px; margin: 40px auto; padding: 0 20px; color: #222; line-height: 1.8; }}
  h1 {{ font-size: 1.6rem; border-bottom: 2px solid #4f8bf9; padding-bottom: 8px; }}
  h2 {{ font-size: 1.2rem; color: #4f8bf9; margin-top: 2rem; }}
  .meta {{ color: #888; font-size: 0.9rem; margin-bottom: 1.5rem; }}
  .speech {{ background: #f8f9fa; border-left: 4px solid #4f8bf9; padding: 12px 16px; margin: 12px 0; border-radius: 4px; }}
  .agent-name {{ font-weight: bold; margin-bottom: 6px; color: #333; }}
  .user-speech {{ background: #fff8e1; border-left: 4px solid #ffc107; padding: 12px 16px; margin: 12px 0; border-radius: 4px; }}
  .summary {{ background: #f0f7ff; border-left: 4px solid #28a745; padding: 16px; border-radius: 4px; }}
  hr {{ border: none; border-top: 1px solid #eee; margin: 1.5rem 0; }}
  @media print {{ body {{ margin: 20px; }} }}
</style>
</head>
<body>
<h1>辩论记录：{st.session_state.topic}</h1>
<div class="meta">
  创建时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} &nbsp;|&nbsp; 辩论轮数：{st.session_state.debate_round}
  {f'&nbsp;|&nbsp; 补充条件：{st.session_state.user_context}' if st.session_state.user_context else ''}
</div>
{''.join(rows)}
{summary_html}
<p style="color:#aaa;font-size:0.8rem;margin-top:2rem;">按 Ctrl+P 可打印或另存为 PDF</p>
</body></html>"""
    return html.encode("utf-8"), f"辩论记录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"

def export_to_markdown() -> tuple:
    if not st.session_state.topic or not st.session_state.debate_history:
        return None, None
    lines = [f"# 辩论记录：{st.session_state.topic}\n",
             f"- 创建时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
             f"- 辩论轮数：{st.session_state.debate_round}"]
    if st.session_state.user_context:
        lines.append(f"- 补充条件：{st.session_state.user_context}")
    lines.append("\n---\n")
    for r in range(1, st.session_state.debate_round + 1):
        lines.append(f"## 第{r}轮\n")
        for item in [x for x in st.session_state.debate_history if x["round"] == r]:
            if item["type"] == "user_speech":
                lines.append(f"**【用户插嘴 → {item.get('target','')}】**\n\n{item['content']}\n")
            else:
                lines.append(f"**{item['name']}**\n\n{item['content']}\n")
        lines.append("---\n")
    if st.session_state.last_summary:
        lines += ["## 决策总结\n", st.session_state.last_summary]
    return "\n".join(lines).encode("utf-8"), f"辩论记录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"

def export_to_word() -> tuple:
    if not st.session_state.topic or not st.session_state.debate_history:
        return None, None
    doc   = Document()
    title = doc.add_heading(f"辩论记录：{st.session_state.topic}", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_heading("一、基本信息", level=1)
    doc.add_paragraph(f"创建时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"辩论轮数：{st.session_state.debate_round}")
    if st.session_state.user_context:
        doc.add_paragraph(f"补充条件：{st.session_state.user_context}")
    doc.add_heading("二、完整辩论记录", level=1)
    for r in range(1, st.session_state.debate_round + 1):
        doc.add_heading(f"第{r}轮", level=2)
        for item in [x for x in st.session_state.debate_history if x["round"] == r]:
            prefix = f"【用户插嘴 → {item.get('target','')}】：" if item["type"] == "user_speech" else f"{item['name']}："
            _add_markdown_paragraph(doc, prefix, item["content"])
            doc.add_paragraph()
    if st.session_state.last_summary:
        doc.add_heading("三、决策总结", level=1)
        _add_markdown_paragraph(doc, "", st.session_state.last_summary)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue(), f"辩论记录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

# ════════════════════════════════════════════════════════════════════
# 意见反馈
# ════════════════════════════════════════════════════════════════════

def save_feedback(content: str):
    feedbacks = load_feedback()
    feedbacks.append({
        "user":    CURRENT_USER,
        "time":    datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "content": content,
        "read":    False,
    })
    gh_save_feedback(feedbacks)

def render_feedback_box():
    with st.expander("💬 意见反馈", expanded=False):
        text = st.text_area("遇到问题或有建议？告诉我们", height=100,
                            placeholder="例如：某个功能有 bug / 希望增加 xxx 功能……",
                            key="feedback_input")
        if st.button("提交反馈", use_container_width=True):
            if text.strip():
                save_feedback(text.strip())
                st.success("感谢反馈，已收到！")
                st.session_state.feedback_input = ""
            else:
                st.warning("请输入反馈内容")

# ════════════════════════════════════════════════════════════════════
# 主页面 UI
# ════════════════════════════════════════════════════════════════════

# 管理员直接显示控制台，不显示辩论界面
if CURRENT_USER == _cfg("ADMIN_USER", "admin"):
    with st.sidebar:
        st.header("⚙️ 设置")
        st.caption(f"👤 {CURRENT_USER}")
        if st.button("登出", use_container_width=True):
            logout()
    render_admin_page()
    st.stop()

st.title("🎯 多智能体辩论助手")
st.divider()

st.session_state.topic = st.text_area(
    "💬 输入你的辩论议题",
    value=st.session_state.topic,
    placeholder="例如：大学生是否应该考研？  /  远程办公利大于弊吗？",
    height=100,
).strip()

st.session_state.user_context = st.text_area(
    "📌 补充条件（可选）",
    value=st.session_state.user_context,
    placeholder="例如：我是大三学生，目标985，家庭条件一般……",
    height=80,
).strip()

# 操作按钮
has_topic   = bool(st.session_state.topic)
has_history = bool(st.session_state.debate_history)
is_debating = st.session_state.is_debating

btn_label  = "🚀 开始辩论" if st.session_state.debate_round == 0 else "🔁 继续辩论"
run_debate = st.button(btn_label,   use_container_width=True, disabled=not has_topic or is_debating, type="primary")
interrupt  = st.button("⚡ 自由辩", use_container_width=True, disabled=not has_history or is_debating)
do_summary = st.button("📊 总结",   use_container_width=True, disabled=not has_history or is_debating)

# 搁置当前辩论，清空状态开新局
if has_history and not is_debating:
    if st.button("📁 搁置并新建辩论", use_container_width=True):
        save_debate_history()
        for k, v in {"topic": "", "user_context": "", "debate_history": [],
                     "debate_round": 0, "current_history_id": None,
                     "last_summary": "", "pending_user_speech": None,
                     "search_results": {}}.items():
            st.session_state[k] = v
        # 重置角色为默认，避免上一个议题的自定义角色污染新辩论
        st.session_state.custom_agents = _DEFAULT_AGENTS
        st.rerun()

if is_debating:
    st.info("⏳ AI 正在生成中，请稍候...")

# 用户插嘴区（有辩论历史时显示）
if has_history:
    st.markdown("""
    <div style="border:2px solid #4f8bf9;border-radius:12px;padding:1rem 1.2rem;margin:0.5rem 0 1rem 0;background:rgba(79,139,249,0.06);">
    <span style="font-size:1.1rem;font-weight:600;">💬 加入讨论</span>
    <span style="font-size:0.85rem;color:#888;margin-left:8px;">输入你的观点，下一轮 AI 会直接回应你</span>
    </div>
    """, unsafe_allow_html=True)
    agent_names = [a["name"] for a in st.session_state.custom_agents]
    target      = st.selectbox("发言对象", ["全体"] + agent_names, key="user_speech_target")
    user_input  = st.text_area("你的观点", height=80,
                               placeholder="例如：我觉得支持派忽略了一个关键问题……",
                               key="user_speech_input")
    if st.button("✅ 提交发言，下一轮 AI 回应我", use_container_width=True, type="primary"):
        if user_input.strip():
            st.session_state.pending_user_speech = {
                "target": target, "content": user_input.strip(),
                "round":  st.session_state.debate_round,
            }
            st.session_state.debate_history.append({
                "type": "user_speech", "round": st.session_state.debate_round,
                "name": "用户", "target": target,
                "content": f"（对{target}）{user_input.strip()}",
            })
            st.success("✅ 已提交，点击「继续辩论」让 AI 回应你")
        else:
            st.warning("请输入内容")
    if st.button("取消发言", use_container_width=True):
        st.session_state.pending_user_speech = None
        st.session_state.debate_history = [
            x for x in st.session_state.debate_history
            if not (x["type"] == "user_speech" and x["round"] == st.session_state.debate_round)
        ]
        st.rerun()

st.divider()

# ════════════════════════════════════════════════════════════════════
# 侧边栏
# ════════════════════════════════════════════════════════════════════

with st.sidebar:
    st.header("⚙️ 设置")
    st.caption(f"👤 {CURRENT_USER}")
    if st.button("登出", use_container_width=True):
        logout()

    render_admin_panel()

    # ── 角色管理 ──────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 🎭 角色管理")
    with st.expander("展开编辑角色"):
        if st.session_state.topic:
            if st.button("✨ 根据议题自动生成角色", use_container_width=True):
                with st.spinner("AI 正在分析议题，生成角色中..."):
                    generated = auto_generate_agents(st.session_state.topic)
                if generated:
                    st.session_state.custom_agents = generated
                    save_agent_config(generated)
                    st.success(f"已生成 {len(generated)} 个角色")
                    st.rerun()
                else:
                    st.error("生成失败，请手动编辑")
        else:
            st.caption("输入议题后可自动生成角色")

        for i, agent in enumerate(st.session_state.custom_agents):
            st.text_input(f"角色{i+1}名", key=f"name_{i}",
                          value=st.session_state.get(f"name_{i}", agent["name"]))
            st.text_area(f"角色{i+1}人设", key=f"prompt_{i}",
                         value=st.session_state.get(f"prompt_{i}", agent["prompt"]), height=80)
            if st.button("🗑️ 删除此角色", key=f"del_{i}"):
                st.session_state.custom_agents.pop(i)
                save_agent_config(st.session_state.custom_agents)
                st.rerun()
            st.markdown("---")

        if st.button("➕ 添加角色", use_container_width=True):
            st.session_state.custom_agents.append({"name": "新角色", "prompt": "请输入人设"})
            save_agent_config(st.session_state.custom_agents)
            st.rerun()
        if st.button("💾 保存角色配置", use_container_width=True):
            for i in range(len(st.session_state.custom_agents)):
                st.session_state.custom_agents[i]["name"]   = st.session_state.get(f"name_{i}", st.session_state.custom_agents[i]["name"])
                st.session_state.custom_agents[i]["prompt"] = st.session_state.get(f"prompt_{i}", st.session_state.custom_agents[i]["prompt"])
            save_agent_config(st.session_state.custom_agents)
            st.success("已保存")

    # ── 自由辩设置 ────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### ⚡ 自由辩设置")
    st.session_state.interrupt_rounds = st.slider(
        "打断次数", min_value=1, max_value=6,
        value=st.session_state.interrupt_rounds
    )

    # ── 学术搜索开关 ──────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 🔍 学术搜索")
    st.session_state.enable_search = st.checkbox(
        "启用学术资料搜索",
        value=st.session_state.enable_search,
        help="每个角色会根据自己的立场，从 Semantic Scholar 和维基百科检索学术资料作为论据支撑"
    )

    # ── 导出 ──────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 📤 导出")
    md_bytes,   md_fn   = export_to_markdown()
    word_bytes, word_fn = export_to_word()
    html_bytes, html_fn = export_to_html()
    if md_bytes:
        st.download_button("导出为 Markdown 文件", data=md_bytes, file_name=md_fn,
                           mime="text/markdown", use_container_width=True)
    else:
        st.button("导出为 Markdown 文件", disabled=True, use_container_width=True)
    if word_bytes:
        st.download_button("导出为 Word 文件", data=word_bytes, file_name=word_fn,
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)
    else:
        st.button("导出为 Word 文件", disabled=True, use_container_width=True)
    if html_bytes:
        st.download_button("导出为 HTML 文件", data=html_bytes, file_name=html_fn,
                           mime="text/html", use_container_width=True)
    else:
        st.button("导出为 HTML 文件", disabled=True, use_container_width=True)
    st.caption("🌐 HTML 用浏览器打开后，Ctrl+P 可转为 PDF")

    # ── 历史记录 ──────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 💾 历史记录")
    if st.button("💾 手动保存当前", use_container_width=True):
        f = save_debate_history()
        if f:
            st.success("已保存")
    if st.button("🧹 清空当前", use_container_width=True):
        for k, v in {"topic": "", "user_context": "", "debate_history": [],
                     "debate_round": 0, "current_history_id": None, "last_summary": ""}.items():
            st.session_state[k] = v
        st.rerun()

    histories = get_all_history()
    if histories:
        labels   = {f: get_history_label(f) for f in histories}
        selected = st.selectbox("选择记录", histories, format_func=lambda f: labels[f])
        hc1, hc2 = st.columns(2)
        with hc1:
            if st.button("🔄 加载", use_container_width=True):
                load_debate_history(selected)
        with hc2:
            if st.button("🗑️ 删除", use_container_width=True):
                delete_history(selected)
    else:
        st.info("暂无历史记录")

# ════════════════════════════════════════════════════════════════════
# 辩论逻辑
# ════════════════════════════════════════════════════════════════════

if run_debate:
    st.session_state.is_debating = True
    st.session_state.debate_round += 1
    r           = st.session_state.debate_round
    is_first    = (r == 1)
    hist_ctx    = format_history_recent(st.session_state.debate_history)
    user_speech = st.session_state.pending_user_speech

    st.subheader(f"📢 第{r}轮辩论")
    if user_speech:
        st.info(f"🙋 用户发言（对象：{user_speech['target']}）：{user_speech['content']}")

    batch = []
    for agent in st.session_state.custom_agents:
        st.markdown(f"### 🗣 {agent['name']}")

        # 为该角色检索学术资料
        if st.session_state.enable_search:
            with st.spinner(f"🔍 {agent['name']} 正在检索学术资料..."):
                refs = search_for_agent(st.session_state.topic, agent["name"], agent["prompt"])
                st.session_state.search_results[f"{r}_{agent['name']}"] = refs
            if refs:
                with st.expander(f"📚 {agent['name']} 的参考来源", expanded=False):
                    st.markdown(refs)

        with st.spinner("思考中..."):
            if user_speech:
                is_targeted = (user_speech["target"] == agent["name"] or user_speech["target"] == "全体")
                prompt = get_user_speech_prompt(
                    agent["prompt"], agent["name"],
                    user_speech["target"], user_speech["content"], is_targeted
                )
            else:
                prompt = get_debate_prompt(agent["prompt"], is_first_round=is_first)
            content = call_llm(prompt, hist_ctx, round_num=r, agent_name=agent["name"])

        st.markdown(content)
        st.divider()
        batch.append({"type": "agent_speech", "round": r, "name": agent["name"], "content": content})

    st.session_state.debate_history.extend(batch)
    st.session_state.pending_user_speech = None
    st.session_state.is_debating = False
    auto_save()
    st.rerun()


if interrupt:
    st.session_state.is_debating = True
    st.session_state.debate_round += 1
    r             = st.session_state.debate_round
    hist_ctx      = format_history_recent(st.session_state.debate_history)
    last_speeches = [x for x in st.session_state.debate_history if x["type"] == "agent_speech"]
    current_target = last_speeches[-1] if last_speeches else None

    st.subheader(f"⚡ 第{r}轮（自由辩·打断/追问）")
    batch = []

    for _ in range(st.session_state.interrupt_rounds):
        if not current_target:
            break
        other_agents = [a for a in st.session_state.custom_agents if a["name"] != current_target["name"]]
        if not other_agents:
            break
        attacker = random.choice(other_agents)
        st.markdown(f"### ⚡ {attacker['name']} 打断 {current_target['name']}")

        if st.session_state.enable_search:
            with st.spinner(f"🔍 {attacker['name']} 正在检索反驳资料..."):
                refs = search_for_agent(st.session_state.topic, attacker["name"], attacker["prompt"])
                st.session_state.search_results[f"{r}_{attacker['name']}"] = refs
            if refs:
                with st.expander(f"📚 {attacker['name']} 的参考来源", expanded=False):
                    st.markdown(refs)

        with st.spinner("组织语言中..."):
            content = call_llm(
                get_interrupt_prompt(attacker["prompt"], current_target["name"], current_target["content"][:300]),
                hist_ctx, round_num=r, agent_name=attacker["name"],
            )
        st.markdown(content)
        st.divider()
        batch.append({"type": "agent_speech", "round": r, "name": attacker["name"],
                      "content": f"（打断{current_target['name']}）{content}"})
        current_target = {"name": attacker["name"], "content": content}

    st.session_state.debate_history.extend(batch)
    st.session_state.is_debating = False
    auto_save()
    st.rerun()


if do_summary:
    st.subheader("📊 最终决策总结")
    with st.spinner("生成总结中..."):
        res = call_llm(_SUMMARIZER_PROMPT, format_history_full(st.session_state.debate_history), max_tokens=800)
    st.session_state.last_summary = res
    st.markdown(res)
    if st.button("🔄 重新生成总结"):
        st.session_state.last_summary = ""
        st.rerun()

elif st.session_state.last_summary:
    st.subheader("📊 最终决策总结")
    st.markdown(st.session_state.last_summary)
    if st.button("🔄 重新生成总结"):
        st.session_state.last_summary = ""
        st.rerun()


# ════════════════════════════════════════════════════════════════════
# 完整辩论记录展示
# ════════════════════════════════════════════════════════════════════

st.divider()
st.subheader("📜 完整辩论记录")
if st.session_state.debate_history:
    for r in range(1, st.session_state.debate_round + 1):
        items = [x for x in st.session_state.debate_history if x["round"] == r]
        with st.expander(f"第{r}轮", expanded=(r == st.session_state.debate_round)):
            for it in items:
                if it["type"] == "user_speech":
                    st.info(f"🙋 用户发言（对象：{it.get('target','全体')}）：{it['content']}")
                else:
                    st.markdown(f"**{it['name']}**")
                    st.markdown(it["content"])
                    st.markdown("---")
else:
    st.info("还未开始辩论，输入议题后点击「🚀 开始辩论」")

st.divider()
render_feedback_box()
